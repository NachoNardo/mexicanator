from docx import Document
from docx.shared import Inches

from elenco import Elenco

class GerarDocx(): 

    document = Document()

    def inserirTítulo(self):
        
        self.document.add_heading('Novela Mexicana', 0)

    def inserirElenco(self, elenco,):
        
        p = self.document.add_paragraph("Elenco completo da novela")
        p = self.document.add_paragraph(elenco)

    def gerar(self,enredo):
        
        p = self.document.add_paragraph(enredo)
        self.document.save('novela.docx')


